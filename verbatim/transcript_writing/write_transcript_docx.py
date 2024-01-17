import logging
from typing import List
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from langcodes import standardize_tag

from ..transcription import Utterance, Transcription
from .write_transcript import WriteTranscript  # Assuming WriteTranscript is in the same directory.

LOG = logging.getLogger(__name__)

def format_seconds(seconds: float) -> str:
    """
    Format seconds into HH:MM:SS.

    Parameters:
        seconds (float): The time in seconds.

    Returns:
        str: The formatted time.
    """
    hours, remainder = divmod(seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f'{int(hours):02}:{int(minutes):02}:{int(seconds):02}'

def short_code_to_bcp47(short_code: str) -> str:
    """
    Convert a language code from a short format to BCP47.

    Parameters:
        short_code (str): The input language code in a short format.

    Returns:
        str: The BCP47 language code.
    """
    try:
        return standardize_tag(short_code)
    except ValueError as e:
        LOG.error(f"Error converting language code: {e}")
        return short_code  # Returning the original code might be a good fallback.



def format_word(paragraph, word, formatting, lang):
    """
    Format a word in a paragraph based on specified formatting options.

    Parameters:
        paragraph (docx.text.paragraph.Paragraph): The paragraph to add the word to.
        word (str): The word text.
        formatting (dict): Formatting options (color, underline, highlight).
        lang (str): The language code of the word.

    Returns:
        None
    """
    run = paragraph.add_run(word)

    if lang:
        #pylint: disable=protected-access
        lang_element = OxmlElement('w:lang')
        language_code = short_code_to_bcp47(lang)
        lang_element.set('{http://www.w3.org/XML/1998/namespace}lang', language_code)
        run._element.append(lang_element)

    if formatting['color'] is not None:
        run.font.color.rgb = RGBColor(*formatting['color'])

    run.underline = formatting['underline']

    if formatting['highlight'] is not None:
        run.font.highlight_color = formatting['highlight']
    else:
        run.font.highlight_color = WD_COLOR_INDEX.AUTO



def write_docx(utterances: List[Utterance], no_timestamps: bool, no_speakers: bool,
               with_confidence: bool, with_language: bool, output_file: str) -> None:
    """
    Write a list of utterances to a Microsoft Word (docx) file.

    Parameters:
        utterances (List[Utterance]): List of Utterance objects to be written.
        no_timestamps (bool): If True, exclude timestamps from the output.
        no_speakers (bool): If True, exclude speaker information from the output.
        with_confidence (bool): If True, include confidence-related formatting.
        with_language (bool): If True, include language information.
        output_file (str): Path to the output Word file.

    Returns:
        None
    """
    doc = Document()

    for utterance in utterances:
        paragraph = doc.add_paragraph()
        header = ""

        if not no_timestamps:
            header += f"[{format_seconds(utterance.start)}]"
        if not no_speakers:
            header += f" {utterance.speaker}"
        if header:
            header += ":"
            run = paragraph.add_run(header)
            run.bold = True

        for word in utterance.words:
            formatting = {'color': None, 'underline': False, 'highlight': None}

            if with_confidence and word.confidence < 0.80:
                formatting['underline'] = True

            if with_confidence and word.confidence < 0.50:
                formatting['highlight'] = WD_COLOR_INDEX.YELLOW

            lang = utterance.language if with_language else None

            format_word(paragraph, word.text, formatting, lang)

    doc.save(output_file)



class WriteTranscriptDocx(WriteTranscript):
    def __init__(self, no_timestamps=False, no_speakers=False, with_confidence=True, with_language=True):
        """
        Constructor for WriteTranscriptDocx class.

        Parameters:
            no_timestamps (bool): If True, exclude timestamps from the output.
            no_speakers (bool): If True, exclude speaker information from the output.
            with_confidence (bool): If True, include confidence-related formatting.
            with_language (bool): If True, include language information.
        """
        self.no_timestamps = no_timestamps
        self.no_speakers = no_speakers
        self.with_confidence = with_confidence
        self.with_language = with_language

    def execute(self, transcription_path: str, output_file: str, **kwargs: dict):
        """
        Execute the transcription writing process to a Microsoft Word (docx) file.

        Parameters:
            transcript (Transcription): The input transcription to be written.
            output_file (str): Path to the output Word file.
            **kwargs: Additional keyword arguments (not used in this implementation).

        Returns:
            None
        """
        transcript:Transcription = Transcription.load(transcription_path)
        output_file = f"{output_file}.docx"
        write_docx(
            utterances=transcript.utterances,
            no_timestamps=self.no_timestamps,
            no_speakers=self.no_speakers,
            with_confidence=self.with_confidence,
            with_language=self.with_language,
            output_file=output_file
        )
