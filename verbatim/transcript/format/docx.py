import logging
from typing import List, Union
from enum import Enum
import numpy as np

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .writer import TranscriptWriterConfig, TimestampStyle, SpeakerStyle, ProbabilityStyle, LanguageStyle, TranscriptWriter
from ..words import VerbatimUtterance, VerbatimWord
from ..formatting import format_milliseconds


LOG = logging.getLogger(__name__)

class Style(Enum):
    NONE = 0
    ITALIC = 1
    BOLD = 2
    UNDERLINE = 3

class DocxParagraph:
    def __init__(self, paragraph:Paragraph):
        self.paragraph = paragraph

    def append(self, text:str, style:Union[Style,List[Style],None] = None):
        run:Run = self.paragraph.add_run(text=text)
        if ((isinstance(style, Style) and Style.BOLD == style) or
            (isinstance(style, list) and Style.BOLD in style)):
            run.bold = True
        if ((isinstance(style, Style) and Style.ITALIC == style) or
            (isinstance(style, list) and Style.ITALIC in style)):
            run.italic = True
        if ((isinstance(style, Style) and Style.UNDERLINE == style) or
            (isinstance(style, list) and Style.UNDERLINE in style)):
            run.underline = True

    def bold(self, text:str):
        self.append(text, Style.BOLD)
    def underline(self, text:str):
        self.append(text, Style.UNDERLINE)
    def italic(self, text:str):
        self.append(text, Style.ITALIC)
class DocxFormatter:
    def __init__(
            self,
            speaker_style:SpeakerStyle,
            timestamp_style:TimestampStyle,
            probability_style:ProbabilityStyle,
            language_style:LanguageStyle
        ):
        self.current_language = None
        self.current_speaker = None
        self.speaker_style = speaker_style
        self.timestamp_style = timestamp_style
        self.probability_style = probability_style
        self.language_style = language_style
        self.last_ts:int = 0

    def _format_timestamp(self, md:DocxParagraph, start_ts:int, end_ts:int):
        if self.timestamp_style == TimestampStyle.none:
            pass

        elif self.timestamp_style == TimestampStyle.minute:
            if self.last_ts is None or start_ts - self.last_ts >= 60 * 16000:
                self.last_ts = ((start_ts / 16000) // 60) * 60 * 16000
                md.bold(f"[{format_milliseconds(start_ts * 1000 / 16000)}]")
                md.append("\n")

        elif self.timestamp_style == TimestampStyle.start:
            md.bold(f"[{format_milliseconds(start_ts * 1000 / 16000)}]:")

        elif self.timestamp_style == TimestampStyle.range:
            md.bold(f"[{format_milliseconds(start_ts * 1000 / 16000)}-{format_milliseconds(end_ts * 1000 / 16000)}]:")

    def _format_speaker(self, md:DocxParagraph, speaker:str):
        if self.speaker_style == SpeakerStyle.none:
            return

        elif self.speaker_style == SpeakerStyle.change:
            if speaker != self.current_speaker:
                self.current_speaker = speaker
                md.bold(f"[{speaker}]")

        elif self.speaker_style == SpeakerStyle.always:
            md.bold(f"[{speaker}]")

    def _format_language(self, md:DocxParagraph, language:str, first_word:bool):
        if self.language_style == LanguageStyle.none:
            pass
        elif self.language_style == LanguageStyle.change:
            if language != self.current_language:
                self.current_language = language
                md.bold(f"[{language}]")
        elif self.language_style == LanguageStyle.always:
            if first_word or language != self.current_language:
                self.current_language = language
                md.bold(f"[{language}]")

    def _format_word_with_probability(self, md:DocxParagraph, word:str, probability:float, utterance_probability:float):
        # pylint: disable=too-many-boolean-expressions
        if ((self.probability_style == ProbabilityStyle.word and probability < 0.90/2) or
            (self.probability_style == ProbabilityStyle.word_75 and probability < 0.75/2) or
            (self.probability_style == ProbabilityStyle.word_50 and probability < 0.50/2) or
            (self.probability_style == ProbabilityStyle.word_25 and probability < 0.25/2) ):
            md.underline(word)
        elif ((self.probability_style == ProbabilityStyle.word and probability < 0.90) or
            (self.probability_style == ProbabilityStyle.word_75 and probability < 0.75) or
            (self.probability_style == ProbabilityStyle.word_50 and probability < 0.50) or
            (self.probability_style == ProbabilityStyle.word_25 and probability < 0.25) ):
            md.italic(word)
        elif ((self.probability_style == ProbabilityStyle.line and utterance_probability < 0.90/2) or
            (self.probability_style == ProbabilityStyle.line_75 and utterance_probability < 0.75/2) or
            (self.probability_style == ProbabilityStyle.line_50 and utterance_probability < 0.50/2) or
            (self.probability_style == ProbabilityStyle.line_25 and utterance_probability < 0.25/2)):
            md.underline(word)
        elif ((self.probability_style == ProbabilityStyle.line and utterance_probability < 0.90) or
            (self.probability_style == ProbabilityStyle.line_75 and utterance_probability < 0.75) or
            (self.probability_style == ProbabilityStyle.line_50 and utterance_probability < 0.50) or
            (self.probability_style == ProbabilityStyle.line_25 and utterance_probability < 0.25)):
            md.italic(word)
        else:
            md.append(word)

    def format_utterance(self, utterance:VerbatimUtterance, out:Document):
        paragraph:Paragraph = out.add_paragraph()
        md:DocxParagraph = DocxParagraph(paragraph=paragraph)
        self._format_timestamp(md=md, start_ts=utterance.start_ts, end_ts=utterance.end_ts)
        self._format_speaker(md=md, speaker=utterance.speaker)

        percentile_25 = np.percentile([w.probability for w in utterance.words], 25)

        #pylint: disable=superfluous-parens
        for i, w in enumerate(utterance.words):
            self._format_language(md=md, language=w.lang, first_word=(i == 0))
            self._format_word_with_probability(md=md, word=w.word, probability=w.probability, utterance_probability=percentile_25)


def write_docx(
        *,
        utterances: List[VerbatimUtterance],
        speaker_style:SpeakerStyle,
        timestamp_style:TimestampStyle,
        probability_style:ProbabilityStyle,
        language_style:LanguageStyle,
        output_file: str) -> None:
    doc = Document()

    formatter:DocxFormatter = DocxFormatter(
            speaker_style=speaker_style,
            timestamp_style=timestamp_style,
            probability_style=probability_style,
            language_style=language_style
            )

    for utterance in utterances:
        formatter.format_utterance(utterance=utterance, out=doc)

    doc.save(output_file)

class DocxTranscriptWriter(TranscriptWriter):
    def __init__(self, config: TranscriptWriterConfig):
        super().__init__(config)
        self.utterances = []
        self.output_file = None

    def open(self, path_no_ext:str):
        self.output_file = f"{path_no_ext}.docx"

    def write(self, utterance:VerbatimUtterance, unacknowledged_utterance:List[VerbatimUtterance] = None, unconfirmed_words:List[VerbatimWord] = None):
        self.utterances.append(utterance)

    def close(self):
        write_docx(
            utterances=self.utterances,
            timestamp_style=self.config.timestamp_style,
            speaker_style=self.config.speaker_style,
            probability_style=self.config.probability_style,
            language_style=self.config.language_style,
            output_file=self.output_file
        )
